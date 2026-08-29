"""
One-off script that generates the small binary fixtures used by the test
suite. Not run automatically — the generated files are committed to the
repo so tests don't need reportlab/Pillow at test time. Re-run manually if
fixtures need to change:

    python tests/fixtures/generate_fixtures.py
"""
from __future__ import annotations

import os

from docx import Document
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

FIXTURES_DIR = os.path.dirname(__file__)


def make_native_pdf():
    path = os.path.join(FIXTURES_DIR, "native.pdf")
    c = canvas.Canvas(path, pagesize=letter)

    # Page 1: heading + paragraphs
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 720, "INTRODUCTION")
    c.setFont("Helvetica", 11)
    c.drawString(72, 690, "This is the first paragraph of native PDF text.")
    c.drawString(72, 670, "It has enough characters to skip OCR entirely.")
    c.showPage()

    # Page 2: another heading + a simple table
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 720, "1. Results")
    c.setFont("Helvetica", 11)
    c.drawString(72, 690, "The table below summarizes results.")

    data = [["Name", "Score"], ["Alice", "90"], ["Bob", "85"]]
    y = 650
    for row in data:
        c.drawString(72, y, f"{row[0]:<10}{row[1]}")
        y -= 18
    # Draw actual grid lines so pdfplumber's table finder detects it.
    x0, y0, col_w, row_h = 72, 660, 100, 18
    for r in range(len(data) + 1):
        c.line(x0, y0 - r * row_h, x0 + 2 * col_w, y0 - r * row_h)
    for cidx in range(3):
        c.line(x0 + cidx * col_w, y0, x0 + cidx * col_w, y0 - len(data) * row_h)
    c.showPage()
    c.save()


def make_scanned_pdf():
    """A PDF with zero native text — an image of text rendered onto each page."""
    path = os.path.join(FIXTURES_DIR, "scanned.pdf")
    img_path = os.path.join(FIXTURES_DIR, "_scanned_page.png")

    image = Image.new("RGB", (1200, 1600), "white")
    draw = ImageDraw.Draw(image)
    draw.text((80, 80), "SCANNED HEADING", fill="black")
    draw.text((80, 160), "This text only exists as a rasterized image,", fill="black")
    draw.text((80, 200), "so it requires OCR to be recovered.", fill="black")
    image.save(img_path)

    c = canvas.Canvas(path, pagesize=letter)
    c.drawImage(img_path, 0, 0, width=letter[0], height=letter[1])
    c.showPage()
    c.save()
    os.remove(img_path)


def make_mixed_pdf():
    """Page 1 = native text, page 2 = scanned image only."""
    path = os.path.join(FIXTURES_DIR, "mixed.pdf")
    img_path = os.path.join(FIXTURES_DIR, "_mixed_page2.png")

    image = Image.new("RGB", (1200, 1600), "white")
    draw = ImageDraw.Draw(image)
    draw.text((80, 80), "SCANNED PAGE TWO", fill="black")
    draw.text((80, 160), "Only recoverable through OCR.", fill="black")
    image.save(img_path)

    c = canvas.Canvas(path, pagesize=letter)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 720, "NATIVE PAGE ONE")
    c.setFont("Helvetica", 11)
    c.drawString(72, 690, "This page has plenty of native, selectable text.")
    c.showPage()
    c.drawImage(img_path, 0, 0, width=letter[0], height=letter[1])
    c.showPage()
    c.save()
    os.remove(img_path)


def make_docx():
    path = os.path.join(FIXTURES_DIR, "sample.docx")
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the first paragraph of the sample document.")
    doc.add_paragraph("This is the second paragraph, right after the first.")
    doc.add_heading("Data", level=2)
    doc.add_paragraph("The table below shows sample data.")
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Score"
    table.rows[1].cells[0].text = "Alice"
    table.rows[1].cells[1].text = "90"
    table.rows[2].cells[0].text = "Bob"
    table.rows[2].cells[1].text = "85"
    doc.save(path)


if __name__ == "__main__":
    make_native_pdf()
    make_scanned_pdf()
    make_mixed_pdf()
    make_docx()
    print("Fixtures written to", FIXTURES_DIR)

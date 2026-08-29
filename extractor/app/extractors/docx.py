"""
DOCX extraction.

Walks the document body in document order (not paragraphs-then-tables,
which is what `python-docx`'s top-level `.paragraphs` / `.tables`
collections give you separately) so headings, paragraphs and tables come
out interleaved the way a reader would encounter them.
"""
from __future__ import annotations

import re

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.extractors.tables import _rows_to_markdown
from app.models import ContentType, ExtractionMethod, ParagraphBlock, TableBlock

_HEADING_RE = re.compile(r"^Heading (\d+)$", re.IGNORECASE)


def _iter_block_items(document: Document):
    """Yield paragraphs and tables in the order they appear in the document body."""
    parent_elm = document.element.body
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def _table_to_block(table: Table, table_index: int) -> TableBlock | None:
    rows = [[cell.text for cell in row.cells] for row in table.rows]
    markdown = _rows_to_markdown(rows)
    if not markdown:
        return None
    return TableBlock(
        markdown=markdown,
        page=None,
        table_index=table_index,
        extraction_method=ExtractionMethod.NATIVE,
        n_rows=len(rows),
        n_cols=max((len(r) for r in rows), default=0),
    )


def extract_docx(file_path: str) -> list[ParagraphBlock | TableBlock]:
    """Extract paragraphs, headings and tables from a DOCX file, in order."""
    document = Document(file_path)
    blocks: list[ParagraphBlock | TableBlock] = []
    table_index = 0

    for item in _iter_block_items(document):
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if not text:
                continue
            heading_match = _HEADING_RE.match(item.style.name or "")
            if heading_match:
                blocks.append(
                    ParagraphBlock(
                        type=ContentType.HEADING,
                        text=text,
                        page=None,
                        extraction_method=ExtractionMethod.NATIVE,
                        heading_level=int(heading_match.group(1)),
                        layout_order=len(blocks),
                    )
                )
            else:
                blocks.append(
                    ParagraphBlock(
                        type=ContentType.PARAGRAPH,
                        text=text,
                        page=None,
                        extraction_method=ExtractionMethod.NATIVE,
                        layout_order=len(blocks),
                    )
                )
        elif isinstance(item, Table):
            table_block = _table_to_block(item, table_index)
            if table_block is not None:
                table_block.layout_order = len(blocks)
                blocks.append(table_block)
                table_index += 1

    return blocks

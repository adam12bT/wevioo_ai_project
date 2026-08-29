"""
Extraction Agent Implementation
-----------------
Runs in PARALLEL with the Research agent — both fan out from the Verifier
and join at Generation, since neither depends on the other (this one reads
the embedded tender doc via RAG; Research goes out to the open web).

Retrieval and generation are split: AnythingLLM's vector-search does pure
similarity search against the embedded tender doc (no LLM call, no
"general knowledge" leakage), and the retrieved chunks are handed as
plain-text context to whichever LLMProvider is configured (Groq or
Ollama — see providers/factory.py) to actually extract the JSON. This
keeps the "answer ONLY from the document" guarantee that AnythingLLM's
mode="query" used to give, without hard-coding which model does the
extraction.
"""

import json
import logging
import os
import re

from .prompts import EXTRACTION_PROMPT
from providers import get_provider
from rfp.default_template import default_response_template

logger = logging.getLogger(__name__)

_RETRIEVAL_QUERY = (
    "all requested work, deliverables, mandatory requirements, domain-specific "
    "constraints, required evidence and forms, submission instructions, eligibility, "
    "contractual and commercial terms, deadlines, budget, evaluation and selection"
)

_TOP_LEVEL_NUMBERED_HEADING = re.compile(
    r"^\s*(?:section\s+)?(?P<number>\d{1,2})[.)]\s+(?P<title>\S.*)$",
    re.IGNORECASE,
)


def _template_paragraphs(file_path: str) -> list[tuple[str, int | None]]:
    """Read template paragraphs locally so RAG cannot omit outline headings."""
    extension = os.path.splitext(file_path)[1].lower()
    paragraphs: list[tuple[str, int | None]] = []

    if extension == ".docx":
        from docx import Document

        document = Document(file_path)
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (paragraph.style.name or "").strip()
            match = re.match(r"heading\s+(\d+)", style_name, re.IGNORECASE)
            level = int(match.group(1)) if match else None
            paragraphs.append((text, level))
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            paragraphs.append((text, None))
        return paragraphs

    if extension == ".pdf":
        from pypdf import PdfReader

        for page in PdfReader(file_path).pages:
            for line in (page.extract_text() or "").splitlines():
                text = line.strip()
                if text:
                    paragraphs.append((text, None))
    return paragraphs


def _extract_template_sections(file_path: str) -> list[str]:
    """Recover the complete ordered outline directly from DOCX/PDF structure."""
    try:
        paragraphs = _template_paragraphs(file_path)
    except Exception as exc:
        logger.warning("Could not inspect response template structure: %s", exc)
        return []

    styled = [(text, level) for text, level in paragraphs if level is not None]
    if styled:
        top_level = min(level for _, level in styled)
        headings = [text for text, level in styled if level == top_level]
        headings = list(dict.fromkeys(headings))
        if len(headings) >= 3:
            return headings

    numbered: list[tuple[int, str]] = []
    seen_numbers = set()
    for text, _ in paragraphs:
        match = _TOP_LEVEL_NUMBERED_HEADING.match(text)
        if not match or len(text) > 180:
            continue
        number = int(match.group("number"))
        if number in seen_numbers:
            continue
        seen_numbers.add(number)
        numbered.append((number, text))

    if len(numbered) < 3:
        return []
    numbers = [number for number, _ in numbered]
    increasing = numbers == sorted(numbers)
    coverage = len(numbers) / max(numbers)
    if increasing and numbers[0] == 1 and coverage >= 0.6:
        return [text for _, text in numbered]
    return []


def _merge_template_outline(requirements: dict, sections: list[str]) -> dict:
    """Prefer a fuller deterministic outline over incomplete RAG extraction."""
    if not sections or not isinstance(requirements, dict):
        return requirements
    template = requirements.get("response_template")
    if not isinstance(template, dict):
        template = {}
        requirements["response_template"] = template
    extracted = template.get("section_order") or template.get("required_sections") or []
    if len(sections) > len(extracted):
        template["required_sections"] = sections
        template["section_order"] = sections
        template["outline_source"] = "local_document_structure"
        logger.info(
            "Recovered %d response-template sections locally (LLM/RAG found %d)",
            len(sections),
            len(extracted),
        )
    return requirements


def _repair_truncated_json(text: str) -> str:
    """Best-effort repair for JSON that was cut off mid-object/array —
    e.g. the LLM's response got truncated by a token limit before it
    could close its braces (this is what happened in the observed run:
    the response ended right after the last complete key/value pair,
    with no closing '}' for the outer object).

    Walks the text respecting string literals/escapes, tracks how many
    '{'/'[' are still open, trims any dangling incomplete token at the
    very end (an unterminated string, a trailing comma, a half-written
    key), then appends the missing closing brackets in the right order.
    """
    stack = []
    in_string = False
    escape = False
    last_safe_index = 0  # index right after the last structurally complete point

    for i, ch in enumerate(text):
        if in_string:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == '"':
                in_string = False
                last_safe_index = i + 1
            continue

        if ch == '"':
            in_string = True
        elif ch in "{[":
            stack.append(ch)
        elif ch in "}]":
            if stack:
                stack.pop()
            last_safe_index = i + 1
        elif ch not in " \t\n\r,":
            last_safe_index = i + 1

    repaired = text[:last_safe_index].rstrip()
    if repaired.endswith(","):
        repaired = repaired[:-1]

    closers = {"{": "}", "[": "]"}
    while stack:
        repaired += closers[stack.pop()]

    return repaired


def _extract_json(text: str) -> dict:
    """LLMs often wrap JSON in prose or markdown fences despite instructions,
    and can also get cut off mid-object by a token limit — handle both."""
    fence_match = re.search(r"```(?:json)?\s*(\{.*\})\s*```", text, re.DOTALL)
    candidate = fence_match.group(1) if fence_match else text

    brace_match = re.search(r"\{.*\}", candidate, re.DOTALL)
    if brace_match:
        candidate = brace_match.group(0)
    else:
        # No closing brace anywhere — very likely truncated. Fall back to
        # everything from the first '{' onward so repair has something to
        # work with, instead of giving up immediately.
        open_brace = candidate.find("{")
        if open_brace != -1:
            candidate = candidate[open_brace:]

    try:
        return json.loads(candidate)
    except json.JSONDecodeError:
        pass

    try:
        repaired = _repair_truncated_json(candidate)
        parsed = json.loads(repaired)
        parsed["_extraction_note"] = (
            "Response appeared truncated (likely hit a token limit); "
            "auto-repaired by closing the open brackets. Spot-check the "
            "last field before trusting it fully."
        )
        logger.warning("Extraction response was truncated; auto-repaired JSON.")
        return parsed
    except json.JSONDecodeError:
        logger.error(
            "Failed to parse extraction response as JSON even after repair "
            "attempt (%d chars of raw response follow at debug level).",
            len(text),
        )
        logger.debug("Unparseable extraction response: %s", text)
        return {"raw_response": text, "parse_error": True}


def extraction_agent(state: dict, *, rag=None) -> dict:
    if not state.get("is_verified"):
        # Should never actually run if the graph is wired correctly, but
        # guard against it anyway rather than silently doing bad work.
        # Partial-return convention: nothing to contribute == empty dict,
        # not a full state passthrough (see state.py docstring — this
        # node runs in parallel with Research, so it must never spread
        # `**state` back).
        return {}

    workspace_slug = state["workspace_slug"]
    template_workspace_slug = state.get("response_template_workspace_slug")
    template_file_path = state.get("response_template_file_path")
    deterministic_sections = (
        _extract_template_sections(template_file_path) if template_file_path else []
    )

    try:
        if rag is None:
            raise RuntimeError("RagQuery dependency was not provided")
        context = rag.query(workspace_slug, _RETRIEVAL_QUERY, top_n=8)
        if template_workspace_slug:
            template_context = rag.query(
                template_workspace_slug,
                "all section headings, exact section order, content instructions, "
                "formatting rules, limits, annexes, forms and mandatory tables",
                top_n=10,
            )
        else:
            template_context = (
                "NO RESPONSE TEMPLATE WAS UPLOADED. Extract tender facts only; "
                "the application will attach its built-in response structure."
            )
        prompt = (
            f"TENDER DOCUMENT EXCERPTS:\n\n{context}\n\n"
            f"RESPONSE TEMPLATE EXCERPTS:\n\n{template_context}\n\n"
            f"{EXTRACTION_PROMPT}"
        )
        response_text = get_provider().complete(
            prompt,
            request_label="extraction.requirements",
            reasoning_effort="low",
            include_reasoning=False,
        )
        requirements = _extract_json(response_text)
        if template_file_path:
            requirements = _merge_template_outline(requirements, deterministic_sections)
            template = requirements.get("response_template")
            if isinstance(template, dict):
                template["template_source"] = "uploaded"
        else:
            requirements["response_template"] = default_response_template()
    except Exception as e:
        error_msg = f"Extraction agent failed: {e}"
        logger.error("Extraction failed for workspace %r: %s", workspace_slug, e, exc_info=True)
        return {
            "requirements": {},
            "errors": [error_msg],
        }

    if requirements.get("parse_error"):
        logger.warning("Extraction produced unparseable requirements for workspace %r", workspace_slug)

    return {"requirements": requirements}

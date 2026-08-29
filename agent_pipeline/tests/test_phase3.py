import unittest
import os
import tempfile
from unittest.mock import Mock, patch

from docx import Document

from rfp.agents.generation.implementation import (
    _batch_output_token_limit,
    _completion_was_truncated,
    _fit_generation_prompt,
    _merge_section_drafts,
    _normalize_batch_headings,
    _proposal_structure,
    _recover_single_section_response,
    _remove_generation_instruction_leaks,
    _search_knowledge_with_trace,
    _has_substantive_section_body,
    _salvage_truncated_section,
    _sanitize_generated_claims,
    _section_requests_company_evidence,
    _section_batches,
    _split_batch_sections,
    generation_agent,
)
from rfp.orchestration.routing import after_generation
from rfp.agents.extraction.implementation import (
    _extract_template_sections,
    _merge_template_outline,
)
from rfp.agents.quality.implementation import (
    _check_section_order,
    _check_template_compliance,
    _evidence_gap_warnings,
    _evaluate_grounding_and_coherence,
    _extract_review_json,
    _identify_failed_sections,
    _insubstantial_sections,
    _reconcile_disclosed_evidence_gaps,
    _reconcile_future_plan_findings,
    _reconcile_tender_requirement_findings,
    _review_groups,
    _template_sections,
    quality_agent,
)


class ResponseTemplateQualityTests(unittest.TestCase):
    def test_explicit_evidence_placeholder_is_a_non_blocking_warning(self):
        sections = ["Qualifications and Supporting Evidence"]
        draft = (
            "## Qualifications and Supporting Evidence\n\n"
            "[TO BE CONFIRMED - supporting company evidence not found]"
        )

        self.assertEqual(_insubstantial_sections(draft, sections), [])
        self.assertEqual(
            _evidence_gap_warnings(draft, sections),
            [
                {
                    "section": sections[0],
                    "placeholder_count": 1,
                    "message": (
                        "Missing supporting evidence. Upload or verify the required source "
                        "documents, then regenerate this section."
                    ),
                }
            ],
        )

    def test_disclosed_evidence_gap_does_not_reduce_quality_scores(self):
        review = {
            "groundedness_score": 0.30,
            "coherence_score": 0.68,
            "unsupported_claims": [
                {
                    "claim": "Team CVs and supporting evidence will be provided [TO BE CONFIRMED].",
                    "reason": "The documents are currently missing.",
                }
            ],
            "contradictions": [],
            "coherence_issues": [
                "The proposal contains placeholders because supporting evidence is missing."
            ],
            "notes": [],
        }

        reconciled = _reconcile_disclosed_evidence_gaps(review)

        self.assertEqual(reconciled["unsupported_claims"], [])
        self.assertEqual(reconciled["coherence_issues"], [])
        self.assertGreaterEqual(reconciled["groundedness_score"], 0.75)
        self.assertGreaterEqual(reconciled["coherence_score"], 0.75)

    def test_confident_unsupported_capability_claim_still_fails(self):
        finding = {
            "claim": "The bidder has verified public-sector delivery experience.",
            "reason": "No company references were supplied.",
        }
        review = {
            "groundedness_score": 0.30,
            "coherence_score": 0.90,
            "unsupported_claims": [finding],
            "contradictions": [],
            "coherence_issues": [],
        }

        reconciled = _reconcile_disclosed_evidence_gaps(review)

        self.assertEqual(reconciled["unsupported_claims"], [finding])
        self.assertEqual(reconciled["groundedness_score"], 0.30)

    def test_generation_guard_neutralizes_unsupported_high_risk_claims(self):
        guarded, findings = _sanitize_generated_claims(
            "The procurement criteria allocate 25% to experience. "
            "The bidder has relevant experience. "
            "The delivery uses a framework aligned with ISO 31000.",
            tender_evidence="A case-management platform is required.",
            company_evidence="(none found in the company knowledge base for this query)",
        )

        self.assertEqual(guarded.count("### Evidence gaps"), 1)
        self.assertIn("Tender-backed commitment", guarded)
        self.assertIn("Company evidence", guarded)
        self.assertNotIn("25%", guarded)
        self.assertNotIn("ISO 31000", guarded)
        self.assertEqual(len(findings), 3)

    def test_generation_guard_neutralizes_missing_cv_and_experience_claims(self):
        guarded, findings = _sanitize_generated_claims(
            "The bidder has relevant experience with similar public-sector systems. "
            "Key personnel possess the required certifications and delivered comparable projects.",
            tender_evidence="The tender requires comparable experience and certified personnel.",
            company_evidence=(
                "(none found in the company knowledge base for this query)\n"
                "(none found in the company knowledge base for this query)"
            ),
        )

        self.assertNotIn("has relevant experience", guarded)
        self.assertNotIn("possess the required certifications", guarded)
        self.assertEqual(guarded.count("### Evidence gaps"), 1)
        self.assertEqual(guarded.count("| Company evidence |"), 1)
        self.assertEqual(len(findings), 2)

    def test_generation_guard_removes_subjectless_experience_claim(self):
        guarded, findings = _sanitize_generated_claims(
            "Demonstrable experience in the design, development, and deployment "
            "of public-sector case management systems is a key strength.",
            tender_evidence="The tender requests a case management system.",
            company_evidence="(none found in the company knowledge base for this query)",
        )

        self.assertNotIn("Demonstrable experience", guarded)
        self.assertEqual(guarded.count("### Evidence gaps"), 1)
        self.assertEqual(guarded.count("| Company evidence |"), 1)
        self.assertEqual(len(findings), 1)

    def test_company_evidence_routing_uses_section_semantics(self):
        self.assertTrue(
            _section_requests_company_evidence(
                ["Qualifications and Supporting Evidence"]
            )
        )
        self.assertTrue(_section_requests_company_evidence(["Équipe proposée"]))
        self.assertFalse(_section_requests_company_evidence(["Executive Summary"]))

    def test_generation_guard_blocks_unverified_artefacts_and_pricing(self):
        guarded, findings = _sanitize_generated_claims(
            "Curriculum Vitae of all personnel have been compiled and are included in Appendix B. "
            "The commercial offer is presented as a lump-sum fee.",
            tender_evidence="The tender requests CVs and a commercial response.",
            company_evidence="(none found in the company knowledge base for this query)",
        )

        self.assertNotIn("Appendix B", guarded)
        self.assertNotIn("lump-sum", guarded)
        self.assertEqual(len(findings), 2)

    def test_quality_keeps_existing_claims_but_ignores_future_project_plans(self):
        review = _reconcile_future_plan_findings(
            {
                "groundedness_score": 0.2,
                "coherence_score": 0.9,
                "unsupported_claims": [
                    {"claim": "A Risk Management Office will be established."},
                    {"claim": "The bidder has verified public-sector experience."},
                ],
                "contradictions": [],
                "coherence_issues": [],
                "notes": [],
            }
        )

        self.assertEqual(len(review["unsupported_claims"]), 1)
        self.assertIn("verified", review["unsupported_claims"][0]["claim"])
        self.assertEqual(len(review["future_plan_findings_ignored"]), 1)

    def test_quality_ignores_historical_proof_demand_for_tender_delivery_plan(self):
        review = _reconcile_future_plan_findings(
            {
                "groundedness_score": 0.3,
                "coherence_score": 0.9,
                "unsupported_claims": [
                    {
                        "claim": "The delivery model follows the six-phase lifecycle defined by the client.",
                        "reason": "No project reference shows that the bidder previously executed this lifecycle.",
                    }
                ],
                "contradictions": [],
                "coherence_issues": ["Minor wording ambiguity."],
                "notes": [],
            }
        )

        self.assertEqual(review["unsupported_claims"], [])
        self.assertGreaterEqual(review["groundedness_score"], 0.75)

    def test_quality_does_not_demand_bidder_history_for_tender_quantities(self):
        state = {
            "generation_evidence": {
                "requirements": (
                    "Integration with six external systems. "
                    "Migration of up to 2 million legacy records."
                )
            }
        }
        review = _reconcile_tender_requirement_findings(
            state,
            {
                "groundedness_score": 0.3,
                "coherence_score": 0.9,
                "unsupported_claims": [
                    {
                        "claim": "Integration with six external systems.",
                        "reason": "No project reference demonstrates integration with six systems.",
                    },
                    {
                        "claim": "Migration of up to 2 million legacy records.",
                        "reason": "No evidence shows the bidder has performed a migration of this scale.",
                    },
                ],
                "contradictions": [],
                "coherence_issues": [],
                "notes": [],
            },
        )

        self.assertEqual(review["unsupported_claims"], [])
        self.assertGreaterEqual(review["groundedness_score"], 0.75)

    def test_company_retrieval_keeps_distinct_source_documents_visible(self):
        knowledge = Mock()
        knowledge.query_with_trace.return_value = {
            "selected": [
                {
                    "chunk_id": "amira-1",
                    "content": "Amira Haddad, Solution Architect",
                    "metadata": {"sourceDocument": "amira-haddad-docx-1.txt"},
                },
                {
                    "chunk_id": "amira-2",
                    "content": "Amira architecture experience",
                    "metadata": {"sourceDocument": "amira-haddad-docx-2.txt"},
                },
                {
                    "chunk_id": "youssef-1",
                    "content": "Youssef Gharbi, Cybersecurity Lead",
                    "metadata": {"sourceDocument": "youssef-gharbi-docx-1.txt"},
                },
                {
                    "chunk_id": "mariem-1",
                    "content": "Mariem Ben Salem, Delivery and Migration Lead",
                    "metadata": {"sourceDocument": "mariem-ben-salem-docx-1.txt"},
                },
            ]
        }

        trace = _search_knowledge_with_trace(
            knowledge, "company-cvs", "all required roles", top_n=12
        )

        self.assertIn("Amira Haddad", trace["context"])
        self.assertIn("Youssef Gharbi", trace["context"])
        self.assertIn("Mariem Ben Salem", trace["context"])
        self.assertEqual(trace["context"].count("source document;"), 3)

    def test_generation_cleanup_removes_internal_word_budget_hint(self):
        cleaned = _remove_generation_instruction_leaks(
            "## Approach\n(360-472 words)\nSubstantive content."
        )

        self.assertNotIn("360-472", cleaned)
        self.assertIn("Substantive content", cleaned)

    def test_single_paraphrased_heading_is_mapped_to_dynamic_section(self):
        expected = "1. Client-defined title / Exact English title"
        body = " ".join(["Grounded section content"] * 12)

        draft, unmatched = _normalize_batch_headings(
            f"## A paraphrased title\n{body}",
            [expected],
        )

        self.assertEqual(unmatched, [])
        self.assertTrue(draft.startswith(f"## {expected}"))
        self.assertIn(body, draft)

    def test_multiple_unexpected_peer_headings_are_not_mapped(self):
        expected = "Dynamic assigned section"

        draft, unmatched = _normalize_batch_headings(
            "## Unexpected one\nContent.\n## Unexpected two\nContent.",
            [expected],
        )

        self.assertEqual(unmatched, [expected])
        self.assertNotIn(f"## {expected}", draft)

    def test_single_section_prose_without_heading_is_preserved(self):
        section = "Dynamic client section"
        prose = " ".join(["Grounded delivery detail"] * 12)

        recovered = _recover_single_section_response(prose, [section], {})

        self.assertEqual(list(recovered), [section])
        self.assertTrue(recovered[section].startswith(f"## {section}"))
        self.assertIn(prose, recovered[section])

    def test_heading_only_response_is_not_treated_as_generated_content(self):
        section = "Dynamic client section"

        recovered = _recover_single_section_response(
            f"## {section}", [section], {}
        )

        self.assertEqual(recovered, {})
        split = _split_batch_sections(f"## {section}", [section])
        self.assertFalse(_has_substantive_section_body(split[section]))

    def test_provider_token_limit_is_detected(self):
        self.assertTrue(_completion_was_truncated({"finish_reason": "length"}))
        self.assertFalse(_completion_was_truncated({"finish_reason": "stop"}))

    def test_truncated_section_is_trimmed_at_complete_sentence(self):
        section = "Dynamic client section"
        complete = " ".join(["Grounded delivery detail"] * 20) + "."
        unfinished = " This final thought was cut in the middle of an unfinished"

        salvaged = _salvage_truncated_section(
            f"## {section}\n\n{complete}{unfinished}",
            maximum_words=200,
        )

        self.assertIn(complete, salvaged)
        self.assertNotIn("unfinished", salvaged)
        self.assertTrue(_has_substantive_section_body(salvaged))

    def test_dynamic_output_budget_has_room_for_largest_section(self):
        with patch.dict(os.environ, {}, clear=False):
            os.environ.pop("GENERATION_BATCH_MAX_TOKENS", None)
            limit = _batch_output_token_limit(
                {"section_order": ["Only section"]},
                1,
            )

        self.assertGreater(limit, 1600)
        self.assertLessEqual(limit, 2400)

    def test_bilingual_short_heading_is_restored_to_exact_template_title(self):
        expected = "5. Plan de travail et calendrier / Work Plan and Timeline"
        draft, unmatched = _normalize_batch_headings(
            "## 5. Plan de travail et calendrier\nContenu détaillé.",
            [expected],
        )

        self.assertIn(f"## {expected}", draft)
        self.assertEqual(unmatched, [])

    def test_numbered_reworded_heading_is_restored_to_exact_template_title(self):
        expected = "10. Proposition financière / Financial Proposal"
        draft, unmatched = _normalize_batch_headings(
            "## 10. Budget et tarification\nMontant fondé sur le cahier des charges.",
            [expected],
        )

        self.assertIn(f"## {expected}", draft)
        self.assertEqual(unmatched, [])

    def test_generated_batch_is_split_into_live_template_sections(self):
        sections = ["1. Introduction", "2. Méthodologie / Methodology"]
        content = _split_batch_sections(
            "## 1. Introduction\nContexte vérifié.\n\n"
            "## 2. Méthodologie / Methodology\nApproche proposée.",
            sections,
        )

        self.assertIn("Contexte vérifié.", content[sections[0]])
        self.assertIn("Approche proposée.", content[sections[1]])

    def test_quality_maps_a_claim_to_only_its_template_section(self):
        sections = ["1. Context", "2. Solution", "3. Schedule"]
        claim = "The proposed platform guarantees unlimited availability."
        draft = (
            "## 1. Context\nTender facts.\n\n"
            f"## 2. Solution\n{claim}\n\n"
            "## 3. Schedule\nThirty-two week schedule."
        )

        failed = _identify_failed_sections(
            draft=draft,
            sections=sections,
            missing_sections=[],
            out_of_order_sections=[],
            quality_findings={},
            grounding_review={
                "groundedness_score": 0.8,
                "coherence_score": 0.9,
                "unsupported_claims": [],
                "contradictions": [{"claim": claim, "evidence": "unsupported"}],
            },
            word_count=200,
        )

        self.assertEqual(failed, ["2. Solution"])

    def test_retry_regenerates_only_failed_section_and_preserves_others(self):
        sections = ["1. Context", "2. Solution", "3. Schedule"]
        previous_draft = (
            "## 1. Context\nAccepted context.\n\n"
            "## 2. Solution\nUnsupported guarantee.\n\n"
            "## 3. Schedule\nAccepted schedule."
        )
        prior_batches = [
            {
                "sections": [section],
                "draft": block,
                "tender_excerpts": "evidence",
                "response_template_excerpts": "instructions",
            }
            for section, block in _split_batch_sections(
                previous_draft, sections
            ).items()
        ]
        rag = Mock()
        rag.query.return_value = "relevant evidence"
        knowledge = Mock()
        provider = Mock()
        provider.complete.return_value = (
            "## 2. Solution\nRepaired solution grounded in tender evidence with concrete "
            "activities, outputs, controls, responsibilities, and acceptance criteria."
        )

        with patch(
            "rfp.agents.generation.implementation.get_provider",
            return_value=provider,
        ):
            result = generation_agent(
                {
                    "is_verified": True,
                    "workspace_slug": "tender",
                    "response_template_workspace_slug": "template",
                    "requirements": {
                        "scope_summary": "digital platform",
                        "response_template": {
                            "required_sections": sections,
                            "section_order": sections,
                        },
                    },
                    "research_summary": "relevant research",
                    "previous_draft": previous_draft,
                    "previous_generation_evidence": {
                        "section_batches": prior_batches,
                        "project_references": "references",
                        "cv_excerpts": "profiles",
                        "past_proposals": "proposals",
                    },
                    "generation_attempts": 1,
                    "quality_report": {"failed_sections": ["2. Solution"]},
                },
                rag=rag,
                knowledge=knowledge,
            )

        self.assertEqual(provider.complete.call_count, 1)
        self.assertIn("Accepted context.", result["draft_proposal"])
        self.assertIn("Repaired solution grounded", result["draft_proposal"])
        self.assertIn("Accepted schedule.", result["draft_proposal"])
        self.assertNotIn("Unsupported guarantee.", result["draft_proposal"])
        self.assertTrue(result["generation_evidence"]["repair_mode"])
        self.assertEqual(
            result["generation_evidence"]["repaired_sections"],
            ["2. Solution"],
        )

    def test_quality_evaluator_falls_back_when_groq_rejects_json_mode(self):
        provider = Mock()
        provider.complete.side_effect = [
            RuntimeError(
                "Groq completion failed: HTTP 400: "
                '{"error":{"code":"json_validate_failed"}}'
            ),
            """{
              "groundedness_score": 0.84,
              "coherence_score": 0.88,
              "unsupported_claims": [],
              "contradictions": [],
              "coherence_issues": [],
              "notes": []
            }""",
        ]

        with patch("rfp.agents.quality.implementation.get_provider", return_value=provider):
            review = _evaluate_grounding_and_coherence(
                {"generation_evidence": {"requirements": {"scope": "Test"}}},
                "# Proposal\nGrounded draft content.",
            )

        self.assertEqual(provider.complete.call_count, 2)
        self.assertEqual(review["groundedness_score"], 0.84)
        self.assertEqual(review["coherence_score"], 0.88)
        self.assertNotIn("evaluation_error", review)

    def test_quality_review_repairs_truncated_json_with_trailing_comma(self):
        malformed = """```json
        {
          "groundedness_score": 0.82,
          "coherence_score": 0.91,
          "unsupported_claims": [],
          "contradictions": [],
          "coherence_issues": [],
          "notes": ["Evaluation completed",]
        """

        review = _extract_review_json(malformed)

        self.assertEqual(review["groundedness_score"], 0.82)
        self.assertEqual(review["coherence_score"], 0.91)
        self.assertEqual(review["notes"], ["Evaluation completed"])

    def test_docx_heading_structure_recovers_complete_template_outline(self):
        document = Document()
        expected = [
            "1. Introduction",
            "2. Compréhension du besoin",
            "3. Conformité aux exigences",
            "4. Approche et méthodologie",
            "5. Équipe et organisation",
            "6. Calendrier du projet",
            "7. Budget et tarification",
            "8. Assurance qualité",
            "9. Gestion des risques",
            "10. Acceptation et garantie",
            "11. Annexes",
        ]
        for heading in expected:
            document.add_heading(heading, level=1)
            document.add_paragraph("Instructions for this section.")
        handle, path = tempfile.mkstemp(suffix=".docx")
        os.close(handle)
        try:
            document.save(path)
            self.assertEqual(_extract_template_sections(path), expected)
        finally:
            os.unlink(path)

    def test_complete_local_outline_replaces_partial_rag_outline(self):
        requirements = {
            "response_template": {
                "required_sections": ["2. Need", "3. Compliance"],
                "section_order": ["2. Need", "3. Compliance"],
                "instructions": ["Keep this instruction"],
            }
        }
        complete = ["1. Introduction", "2. Need", "3. Compliance", "4. Annexes"]

        merged = _merge_template_outline(requirements, complete)

        self.assertEqual(merged["response_template"]["section_order"], complete)
        self.assertEqual(
            merged["response_template"]["instructions"],
            ["Keep this instruction"],
        )
        self.assertEqual(
            merged["response_template"]["outline_source"],
            "local_document_structure",
        )

    def test_generation_prompt_enforces_total_budget_without_cutting_instructions(self):
        huge = "French tender evidence and requirements. " * 1000
        prompt, fitted = _fit_generation_prompt(
            {
                "batch_number": 1,
                "batch_count": 2,
                "tender_excerpts": huge,
                "response_template_excerpts": huge,
                "response_template_rules": huge,
                "proposal_structure": "## 1. Introduction\n## 2. Methodology",
                "revision_feedback": huge,
                "requirements": huge,
                "research_summary": huge,
                "project_references": huge,
                "cv_excerpts": huge,
                "past_proposals": huge,
            },
            max_chars=13000,
        )

        self.assertLessEqual(len(prompt), 13000)
        self.assertIn("## 1. Introduction", prompt)
        self.assertIn("Do not invent specific figures", prompt)
        self.assertLess(len(fitted["research_summary"]), len(huge))

    def test_quality_failure_does_not_regenerate_by_default(self):
        sections = ["Contexte", "Solution", "Planning"]
        draft = "\n".join(
            f"# {section}\n" + ("substantive content " * 60)
            for section in sections
        )
        evaluator_result = {
            "groundedness_score": 0.5,
            "coherence_score": 0.9,
            "unsupported_claims": [{"claim": "x", "reason": "unsupported"}],
            "contradictions": [],
            "coherence_issues": [],
            "notes": [],
        }

        with patch(
            "rfp.agents.quality.implementation._evaluate_grounding_and_coherence",
            return_value=evaluator_result,
        ):
            result = quality_agent(
                {
                    "is_verified": True,
                    "security_passed": True,
                    "draft_proposal": draft,
                    "generation_attempts": 1,
                    "requirements": {
                        "response_template": {
                            "required_sections": sections,
                            "section_order": sections,
                        }
                    },
                }
            )

        self.assertFalse(result["quality_passed"])

    def test_quality_review_uses_one_compact_group_by_default(self):
        section_batches = [
            {"sections": [f"S{index}"], "draft": f"draft {index}", "evidence": index}
            for index in range(4)
        ]
        groups = _review_groups(
            {"generation_evidence": {"section_batches": section_batches}},
            "full draft",
        )

        self.assertEqual(len(groups), 1)
        self.assertIn("draft 0", groups[0][1])
        self.assertIn("draft 1", groups[0][1])
        self.assertIn("draft 2", groups[0][1])
        self.assertNotIn("draft", groups[0][0]["section_evidence"][0])

    def test_quality_review_prefers_exact_batch_company_evidence(self):
        section_batches = [
            {
                "sections": ["Proposed Team"],
                "draft": "Alice is the proposed security architect.",
                "project_references": "batch project reference",
                "cv_excerpts": "Alice is a certified security architect.",
                "past_proposals": "batch proposal example",
                "requirements": "Security architect required.",
                "research_summary": "external market context",
                "tender_excerpts": "Tender requires a security architect.",
                "response_template_excerpts": "Describe the proposed team.",
            }
        ]
        groups = _review_groups(
            {
                "generation_evidence": {
                    "section_batches": section_batches,
                    "cv_excerpts": "stale top-level CV",
                    "project_references": "stale top-level reference",
                }
            },
            "full draft",
        )

        company = groups[0][0]["company_knowledge"]
        self.assertIn("Alice", company["cv_excerpts"])
        self.assertIn("batch project", company["project_references"])
        self.assertNotIn("stale", company["cv_excerpts"])
        self.assertEqual(
            groups[0][0]["market_research"]["summary"],
            "external market context",
        )

    def test_quality_review_keeps_one_balanced_group_for_large_draft(self):
        section_batches = [
            {
                "sections": [f"S{index}"],
                "draft": f"section-{index} " + ("detail " * 650),
            }
            for index in range(4)
        ]

        groups = _review_groups(
            {"generation_evidence": {"section_batches": section_batches}},
            "full draft",
        )

        self.assertEqual(len(groups), 1)
        reviewed = groups[0][1]
        for index in range(4):
            self.assertIn(f"section-{index}", reviewed)
        self.assertLessEqual(len(reviewed), 6000)

    def test_quality_evaluator_error_does_not_retry_generation(self):
        draft = "\n".join(
            f"# {section}\n" + ("substantive content " * 30)
            for section in [
                "Executive Summary",
                "Understanding of the Requirements",
                "Proposed Approach & Methodology",
                "Indicative Work Plan / Timeline",
                "Risk Management & Quality Assurance",
                "Proposed Team (Profils Proposés)",
                "Why Us",
            ]
        )
        evaluator_result = {
            "groundedness_score": 0.0,
            "coherence_score": 0.0,
            "unsupported_claims": [],
            "contradictions": [],
            "coherence_issues": [],
            "notes": [],
            "evaluation_error": "HTTP 413 request too large",
        }

        with patch(
            "rfp.agents.quality.implementation._evaluate_grounding_and_coherence",
            return_value=evaluator_result,
        ):
            result = quality_agent(
                {
                    "is_verified": True,
                    "security_passed": True,
                    "draft_proposal": draft,
                    "generation_attempts": 1,
                }
            )

        self.assertFalse(result["quality_passed"])
        self.assertFalse(result["quality_passed"])
        self.assertFalse(result["quality_report"]["evaluation_available"])
        self.assertFalse(
            any(
                "groundedness=0.00" in note
                for note in result["quality_report"]["notes"]
            )
        )
        self.assertTrue(
            any("without regenerating" in note for note in result["quality_report"]["notes"])
        )

    def test_template_sections_are_batched_dynamically_in_groups_of_three(self):
        sections = [f"Custom section {index}" for index in range(1, 8)]

        batches = _section_batches(
            {"section_order": sections},
            batch_size=3,
        )

        self.assertEqual(
            batches,
            [sections[0:3], sections[3:6], sections[6:7]],
        )

    def test_empty_generation_stops_before_security(self):
        self.assertNotEqual(after_generation({"generation": {"draft_proposal": ""}}), "security")
        self.assertEqual(
            after_generation({"generation": {"draft_proposal": "proposal"}}), "security"
        )

    def test_client_template_replaces_default_generation_outline(self):
        outline = _proposal_structure(
            {
                "required_sections": ["1. Introduction", "2. Compréhension du besoin"],
                "section_order": ["1. Introduction", "2. Compréhension du besoin"],
            }
        )

        self.assertIn("## 1. Introduction", outline)
        self.assertIn("## 2. Compréhension du besoin", outline)
        self.assertNotIn("Executive Summary", outline)

    def test_generation_outline_derives_word_target_from_template_limit(self):
        short_outline = _proposal_structure(
            {
                "section_order": ["Introduction", "Solution"],
                "formatting_requirements": ["Maximum 10 pages."],
            }
        )
        long_outline = _proposal_structure(
            {
                "section_order": [
                    "Introduction",
                    "Understanding",
                    "Solution",
                    "Planning",
                    "Quality",
                ],
                "formatting_requirements": ["Maximum 10 pages."],
            }
        )

        self.assertIn("derived from template page limit (10 pages)", short_outline)
        self.assertIn("Target length: 650-750 words.", short_outline)
        self.assertIn("Target length: 392-514 words.", long_outline)

    def test_client_template_sections_override_defaults(self):
        state = {
            "requirements": {
                "response_template": {
                    "required_sections": ["Contexte", "Solution", "Planning"],
                    "section_order": ["Contexte", "Solution", "Planning"],
                }
            }
        }

        required, ordered = _template_sections(state)

        self.assertEqual(required, ["Contexte", "Solution", "Planning"])
        self.assertEqual(ordered, required)

    def test_missing_and_out_of_order_sections_are_reported(self):
        draft = "# Solution\nDetails\n# Contexte\nDetails"
        sections = ["Contexte", "Solution", "Planning"]

        self.assertEqual(_check_template_compliance(draft, sections), ["Planning"])
        self.assertEqual(_check_section_order(draft, sections), sections)

    def test_numbered_template_title_matches_unnumbered_markdown_heading(self):
        draft = "# Introduction\nTexte\n## **2. Compréhension du besoin**\nTexte"
        sections = ["1. Introduction", "2. Compréhension du besoin"]

        self.assertEqual(_check_template_compliance(draft, sections), [])
        self.assertEqual(_check_section_order(draft, sections), [])


if __name__ == "__main__":
    unittest.main()
